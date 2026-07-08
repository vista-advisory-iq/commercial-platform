"""
Management Investment Committee report — a single dossier of everything that has
happened on a deal, for the committee meeting where the Stage 2 GO / NO-GO
decision is taken.

Two renderers off one gathered data structure:
  * build_pdf  — reportlab (already in the stack, used for proposal PDFs)
  * build_docx — python-docx, so the committee can annotate in Word

Read-only: this only reads the deal, its scoring and its audit trail. It never
mutates state (all state changes go through services.py).
"""
from io import BytesIO

from .models import DealState
from . import scoring as scoring_engine


# --- Shared data gathering ---------------------------------------------------

_STAGE2_LABELS = {"GO": "GO", "CONDITIONAL": "Conditional GO", "NO_GO": "NO-GO"}


def _fmt(value, suffix=""):
    if value is None or value == "":
        return "—"
    return f"{value}{suffix}"


def _intake_rows(intake):
    """Curated label/value pairs for the intake, grouped for the report."""
    if intake is None:
        return [], []
    identity = [
        ("Deal name", intake.deal_name),
        ("Deal type", intake.get_deal_type_display() if intake.deal_type else ""),
        ("Sub-sector", intake.sub_sector),
        ("Client", intake.client_name),
        ("Counterparty class", intake.counterparty_class),
        ("Location", intake.location),
        ("Sponsor", intake.sponsor),
        ("Sponsor track record (yrs)", intake.sponsor_years),
        ("Deal source", intake.deal_source),
    ]
    economics = [
        ("Total project cost (USD m)", intake.total_project_cost_usd_m),
        ("Installed capacity (kWh)", intake.installed_capacity),
        ("Proposed tariff (NGN/kWh)", intake.proposed_tariff_ngn_kwh),
        ("Tenor (years)", intake.tenor_years),
        ("Revenue in 2–3 yrs (%)", intake.revenue_2_3yr_pct),
        ("Capital structure", intake.capital_structure),
        ("EBITDA (USD m)", intake.ebitda_usd_m),
        ("Leverage (USD m)", intake.leverage_usd_m),
        ("Cash position (USD m)", intake.cash_position_usd_m),
    ]
    return identity, economics


def _gate_rows(deal):
    """Each of the five gates with the recorded verdict."""
    from apps.screening.models import KnockoutGate
    results = {r.gate_id: r for r in deal.gate_results.all()}
    rows = []
    for g in KnockoutGate.objects.all():
        r = results.get(g.id)
        rows.append((
            f"Gate {g.number}: {g.name}",
            (r.verdict if r else "Not assessed"),
            (r.evidence_notes if r else ""),
        ))
    return rows


def gather(deal):
    """Everything the report needs, computed once, as plain Python."""
    intake = getattr(deal, "intake", None)
    identity, economics = _intake_rows(intake)
    scores = scoring_engine.compute_scores(deal)

    decision = None
    if deal.state in (DealState.STAGE2_GO, DealState.STAGE2_CONDITIONAL, DealState.STAGE2_NO_GO):
        decision = {
            "label": _STAGE2_LABELS.get(deal.stage2_decision, deal.stage2_decision or "—"),
            "rationale": deal.stage2_rationale,
        }

    recommendation = _STAGE2_LABELS.get(scores.get("verdict"), scores.get("verdict"))
    history = list(
        deal.state_history.select_related("actor").all()
    )
    return {
        "deal": deal,
        "title": (intake.deal_name if intake and intake.deal_name else deal.deal_ref),
        "identity": identity,
        "economics": economics,
        "gates": _gate_rows(deal),
        "stage1_decision": deal.get_stage1_decision_display() if deal.stage1_decision else "—",
        "scores": scores,
        "recommendation": recommendation or "Not banded",
        "decision": decision,
        "history": history,
    }


# --- PDF (reportlab) ---------------------------------------------------------

def build_pdf(deal):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )

    data = gather(deal)
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=f"MIC Report — {deal.deal_ref}",
    )
    styles = getSampleStyleSheet()
    brand = colors.HexColor("#1e3a5f")
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=18, spaceAfter=2, textColor=brand)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=12.5,
                        spaceBefore=12, spaceAfter=4, textColor=brand)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=9.5, leading=13)
    meta = ParagraphStyle("meta", parent=styles["Normal"], fontSize=9, textColor=colors.grey)
    el = []

    def kv_table(rows):
        clean = [[Paragraph(str(k), body), Paragraph(_fmt(v).replace("\n", "<br/>"), body)]
                 for k, v in rows]
        t = Table(clean, colWidths=[62 * mm, 102 * mm])
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#dfe3e8")),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f4f6f8")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return t

    # Header
    el.append(Paragraph("Management Investment Committee — Deal Report", h1))
    el.append(Paragraph(
        f"{deal.deal_ref} · {data['title']} · Current status: {deal.get_state_display()}",
        meta,
    ))
    el.append(Spacer(1, 6))

    # Recommendation / decision banner
    banner_rows = [
        ["Weighted score", _fmt(data["scores"].get("weighted_total"), " / 100")],
        ["Computed recommendation", data["recommendation"]],
        ["Stage 1 outcome", data["stage1_decision"]],
    ]
    if data["decision"]:
        banner_rows.append(["Committee decision", data["decision"]["label"]])
    el.append(kv_table(banner_rows))
    if data["decision"] and data["decision"]["rationale"]:
        el.append(Spacer(1, 4))
        el.append(Paragraph(f"<b>Decision rationale:</b> {data['decision']['rationale']}", body))

    # Intake
    el.append(Paragraph("Project Identity", h2))
    el.append(kv_table(data["identity"]))
    el.append(Paragraph("Economics", h2))
    el.append(kv_table(data["economics"]))

    # Stage 1
    el.append(Paragraph("Stage 1 — Knockout Gates", h2))
    grows = [[Paragraph("<b>Gate</b>", body), Paragraph("<b>Verdict</b>", body),
              Paragraph("<b>Evidence</b>", body)]]
    for name, verdict, notes in data["gates"]:
        grows.append([Paragraph(name, body), Paragraph(_fmt(verdict), body),
                      Paragraph(_fmt(notes).replace("\n", "<br/>"), body)])
    gt = Table(grows, colWidths=[64 * mm, 26 * mm, 74 * mm])
    gt.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#dfe3e8")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef1f4")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    el.append(gt)

    # Stage 2 scoring
    el.append(Paragraph("Stage 2 — Pillar Scoring", h2))
    for p in data["scores"]["pillars"]:
        el.append(Paragraph(
            f"<b>{p['code']}. {p['name']}</b> — weight {p['weight_pct']}% · "
            f"raw {p['raw']}/{p['max']} · {p['pct']}%",
            body,
        ))
        srows = [[Paragraph("<b>Sub-criterion</b>", body), Paragraph("<b>Figure</b>", body),
                  Paragraph("<b>Score</b>", body)]]
        for sc in p["sub_criteria"]:
            figure = sc.get("measured_value")
            if figure is not None and sc.get("unit"):
                figure = f"{figure} {sc['unit']}"
            srows.append([
                Paragraph(sc["name"], body),
                Paragraph(_fmt(figure), body),
                Paragraph(_fmt(sc["grade"]), body),
            ])
        st = Table(srows, colWidths=[104 * mm, 40 * mm, 20 * mm])
        st.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#e7eaee")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f4f6f8")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        el.append(st)
        if p.get("comment"):
            el.append(Paragraph(f"<i>{p['comment']}</i>", body))
        el.append(Spacer(1, 4))

    # Audit trail
    el.append(Paragraph("Audit Trail — State History", h2))
    for h in data["history"]:
        who = (h.actor.full_name or h.actor.email) if h.actor else "system"
        line = (f"{h.occurred_at:%Y-%m-%d %H:%M} · {_fmt(h.from_state)} → {h.to_state} · {who}")
        if h.reason:
            line += f" — {h.reason}"
        el.append(Paragraph(line, meta))

    doc.build(el)
    buffer.seek(0)
    return buffer.getvalue()


# --- DOCX (python-docx) ------------------------------------------------------

def build_docx(deal):
    from docx import Document
    from docx.shared import Pt, RGBColor

    data = gather(deal)
    document = Document()
    brand = RGBColor(0x1E, 0x3A, 0x5F)

    def heading(text, size=13):
        p = document.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = brand
        return p

    def kv_table(rows):
        table = document.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for k, v in rows:
            cells = table.add_row().cells
            cells[0].text = str(k)
            cells[1].text = _fmt(v)
        document.add_paragraph()

    title = document.add_paragraph()
    trun = title.add_run("Management Investment Committee — Deal Report")
    trun.bold = True
    trun.font.size = Pt(18)
    trun.font.color.rgb = brand
    sub = document.add_paragraph(
        f"{deal.deal_ref} · {data['title']} · Current status: {deal.get_state_display()}"
    )
    sub.runs[0].font.size = Pt(9)

    banner = [
        ("Weighted score", _fmt(data["scores"].get("weighted_total"), " / 100")),
        ("Computed recommendation", data["recommendation"]),
        ("Stage 1 outcome", data["stage1_decision"]),
    ]
    if data["decision"]:
        banner.append(("Committee decision", data["decision"]["label"]))
    kv_table(banner)
    if data["decision"] and data["decision"]["rationale"]:
        document.add_paragraph(f"Decision rationale: {data['decision']['rationale']}")

    heading("Project Identity")
    kv_table(data["identity"])
    heading("Economics")
    kv_table(data["economics"])

    heading("Stage 1 — Knockout Gates")
    gt = document.add_table(rows=1, cols=3)
    gt.style = "Light Grid Accent 1"
    hdr = gt.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Gate", "Verdict", "Evidence"
    for name, verdict, notes in data["gates"]:
        c = gt.add_row().cells
        c[0].text, c[1].text, c[2].text = name, _fmt(verdict), _fmt(notes)
    document.add_paragraph()

    heading("Stage 2 — Pillar Scoring")
    for p in data["scores"]["pillars"]:
        document.add_paragraph(
            f"{p['code']}. {p['name']} — weight {p['weight_pct']}% · "
            f"raw {p['raw']}/{p['max']} · {p['pct']}%"
        )
        st = document.add_table(rows=1, cols=3)
        st.style = "Light Grid Accent 1"
        sh = st.rows[0].cells
        sh[0].text, sh[1].text, sh[2].text = "Sub-criterion", "Figure", "Score"
        for sc in p["sub_criteria"]:
            figure = sc.get("measured_value")
            if figure is not None and sc.get("unit"):
                figure = f"{figure} {sc['unit']}"
            c = st.add_row().cells
            c[0].text, c[1].text, c[2].text = sc["name"], _fmt(figure), _fmt(sc["grade"])
        if p.get("comment"):
            document.add_paragraph(p["comment"]).italic = True
        document.add_paragraph()

    heading("Audit Trail — State History")
    for h in data["history"]:
        who = (h.actor.full_name or h.actor.email) if h.actor else "system"
        line = f"{h.occurred_at:%Y-%m-%d %H:%M} · {_fmt(h.from_state)} → {h.to_state} · {who}"
        if h.reason:
            line += f" — {h.reason}"
        document.add_paragraph(line).runs[0].font.size = Pt(9)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# --- Knockout screening memo (PDF) --------------------------------------------

def _pdf_shell(title_text, deal, subtitle):
    """Common reportlab scaffolding for the standalone screening documents."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=f"{title_text} — {deal.deal_ref}",
    )
    styles = getSampleStyleSheet()
    brand = colors.HexColor("#1e3a5f")
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=17, spaceAfter=2, textColor=brand)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=12.5,
                        spaceBefore=12, spaceAfter=4, textColor=brand)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=10, leading=14)
    meta = ParagraphStyle("meta", parent=styles["Normal"], fontSize=9, textColor=colors.grey)
    el = [Paragraph(title_text, h1), Paragraph(subtitle, meta), Spacer(1, 8)]
    return buffer, doc, el, (h2, body, meta)


def build_screening_memo_pdf(deal):
    """
    Stage 1 knockout screening memo, modelled on the MTN Data Centre memo: the
    per-gate verdicts with evidence, the overall outcome, and any conditions.
    """
    from django.utils import timezone
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import mm

    data = gather(deal)
    buffer, doc, el, (h2, body, meta) = _pdf_shell(
        "Knockout Screening Memo", deal,
        f"{deal.deal_ref} · {data['title']} · Prepared {timezone.now():%d %B %Y} · "
        f"Status: {deal.get_state_display()}",
    )

    intake = getattr(deal, "intake", None)
    summary_rows = [
        ("Client", getattr(intake, "client_name", "") or "—"),
        ("Location", getattr(intake, "location", "") or "—"),
        ("Deal source", getattr(intake, "deal_source", "") or "—"),
        ("Stage 1 outcome", data["stage1_decision"]),
    ]
    agent = getattr(intake, "sales_agent", None)
    if agent:
        summary_rows.insert(3, ("Originating sales agent", str(agent)))
    srows = [[Paragraph(f"<b>{k}</b>", body), Paragraph(_fmt(v), body)] for k, v in summary_rows]
    st = Table(srows, colWidths=[52 * mm, 112 * mm])
    st.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#dfe3e8")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f4f6f8")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    el.append(st)

    el.append(Paragraph("Gate-by-Gate Assessment", h2))
    grows = [[Paragraph("<b>Gate</b>", body), Paragraph("<b>Verdict</b>", body),
              Paragraph("<b>Evidence / notes</b>", body)]]
    for name, verdict, notes in data["gates"]:
        grows.append([Paragraph(name, body), Paragraph(_fmt(verdict), body),
                      Paragraph(_fmt(notes).replace("\n", "<br/>"), body)])
    gt = Table(grows, colWidths=[60 * mm, 26 * mm, 78 * mm])
    gt.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#dfe3e8")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef1f4")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    el.append(gt)

    # Outcome narrative: the most recent state-history reason usually carries
    # the derived Stage 1 rationale (which gates failed / conditions attached).
    el.append(Paragraph("Outcome", h2))
    stage1_reasons = [
        h.reason for h in data["history"]
        if h.reason and ("Stage 1" in h.reason or "Knockout" in h.reason or "gates" in h.reason)
    ]
    if stage1_reasons:
        el.append(Paragraph(stage1_reasons[-1], body))
    else:
        el.append(Paragraph(f"Stage 1 outcome: {data['stage1_decision']}.", body))

    el.append(Spacer(1, 10))
    el.append(Paragraph(
        "This memo is the standard output of every knockout screening. Verdicts and "
        "evidence above are recorded in the deal's append-only audit trail.", meta,
    ))
    doc.build(el)
    buffer.seek(0)
    return buffer.getvalue()


# --- Decline letter (PDF) ------------------------------------------------------

def build_decline_letter_pdf(deal):
    """
    Formal decline letter to the client, generated from the logged decline
    rationale (DECLINED at Stage 1 or NO-GO at Stage 2). The framework requires
    every decline to carry a reasoned letter — this makes it a one-click output.
    """
    from django.utils import timezone
    from reportlab.platypus import Paragraph, Spacer

    data = gather(deal)
    intake = getattr(deal, "intake", None)
    client = getattr(intake, "client_name", "") or "the counterparty"
    contact = getattr(intake, "primary_contact_name", "") or ""

    # The decline reason lives in the state-history row that moved the deal to
    # its declined state.
    declined_row = next(
        (h for h in reversed(data["history"])
         if h.to_state in (DealState.DECLINED, DealState.STAGE2_NO_GO)),
        None,
    )
    reason = (declined_row.reason if declined_row and declined_row.reason else
              "the opportunity does not meet our current investment criteria")

    buffer, doc, el, (h2, body, meta) = _pdf_shell(
        "Re: Outcome of Commercial Screening", deal,
        f"{deal.deal_ref} · {data['title']} · {timezone.now():%d %B %Y}",
    )

    paragraphs = [
        f"Dear {contact or client},",
        f"Thank you for the opportunity to evaluate <b>{data['title']}</b>. "
        "We have completed our commercial screening of the proposal.",
        "After careful assessment against our screening framework, we regret to "
        "advise that we will not be progressing this opportunity at this time. "
        f"In summary: {reason}",
        "This decision reflects our screening criteria as they stand today. Should "
        "the underlying circumstances change materially, we would welcome the "
        "opportunity to revisit the proposal.",
        "We appreciate the time and information you shared with us during the "
        "process, and we hope to find an opportunity to work together in the future.",
        "Yours sincerely,",
        "Commercial Department",
    ]
    for p in paragraphs:
        el.append(Paragraph(p, body))
        el.append(Spacer(1, 8))

    el.append(Spacer(1, 6))
    el.append(Paragraph(
        "Generated from the deal's audit record; the full decline rationale is "
        "retained internally in the append-only history.", meta,
    ))
    doc.build(el)
    buffer.seek(0)
    return buffer.getvalue()
